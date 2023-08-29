from docx import Document
from docx.oxml.ns import qn
from PIL import Image
from io import BytesIO
from image_or_tables_identification import has_tables_or_images
import os
import pandas as pd


def extract_tables(doc):
    tables = []

    for table in doc.tables:
        table_content = []

        for row in table.rows:
            row_content = []
            for cell in row.cells:
                cell_text = ""
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text + " "
                row_content.append(cell_text.strip())
            
            table_content.append(row_content)
        
        tables.append(table_content)
    
    return tables

def extract_images(doc):
    images = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            for run in element.findall('.//' + qn('w:drawing')):
                image_data = run.find('.//' + qn('a:blip')).get(qn('r:embed'))
                if image_data:
                    image_part = doc.part.related_parts[image_data]
                    image_stream = io.BytesIO(image_part.blob)
                    image = Image.open(image_stream)
                    images.append(image)
    
    return images

def save_images(images, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    for image_index, image in enumerate(images, start=1):
        image_filename = os.path.join(output_folder, f"image_{image_index}.png")
        image.save(image_filename, format="PNG")
        print(f"Image {image_index} saved as {image_filename}")


def create_docx_with_tables(tables, output_path):
    doc = Document()

    for table_index, table_content in enumerate(tables, start=1):
        table = doc.add_table(rows=len(table_content), cols=max(len(row) for row in table_content))
        
        for row_index, row_content in enumerate(table_content):
            row_cells = table.rows[row_index].cells
            for col_index, cell_content in enumerate(row_content):
                row_cells[col_index].text = cell_content
    
        if table_index < len(tables):
            # Add a page break after each table except the last one
            doc.add_page_break()

    doc.save(output_path)
    print(f"DOCX file with tables saved as {output_path}")

# ...

def main():
    docx_file_path = 'Etude de cas 2023 - processus commercial Siham hannouda .docx'
    output_folder = 'output_images'  # Folder where PNG images will be saved
    doc = Document(docx_file_path)
    
    has_tables, has_images = has_tables_or_images(doc)
    # print(has_images, has_tables)
    # if has_tables:
    #     tables = extract_tables(doc)
    #     for table_index, table_content in enumerate(tables, start=1):
    #         print(f"Table {table_index}:")
    #         for row_index, row_content in enumerate(table_content, start=1):
    #             print(f"Row {row_index}: {row_content}")
    #         print("-" * 20)
        
    #     output_docx_path = 'output_tables.docx'
    #     create_docx_with_tables(tables, output_docx_path)

    if has_images:
        images = extract_images(doc)
        save_images(images, output_folder)  # Save images as PNG files in the output folder

if __name__ == "__main__":
    main()
