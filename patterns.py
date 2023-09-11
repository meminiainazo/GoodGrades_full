import pandas as pd
#from fill_variables import *
import docx
import re

#Get number of questions ----------------------------------------------------------------------------------------
def get_nb_questions():
    not_Check = True
    while not_Check:
        nb_Questions = input("Entrez le nombre de questions : ")
        try:
            nb_Questions = int(nb_Questions)
            not_Check = False
        except ValueError:
            print("Veuillez entrer un nombre entier !")
    return nb_Questions
# --------------------------------------------------------------------------------------------------------------


# Get patterns  ------------------------------------------------------------------------------------------
def get_patterns(subject):
    pattern_Debut = []
    pattern_End = []
    pattern_point = [r"points\)",r"point\)",r"pts\)",r"pt\)"]
    pattern_numero = [[f"{num}\)",f"{num}/",f"Q{num}",f"QUESTION {num}",f"QUESTION{num}"] for num in range(1,10)]

    if subject.endswith(".docx"): #only for docx, pdf file scheduling
        file = docx.Document(subject)
        numero_is_inside = False
        for paragraph in file.paragraphs:
            for i in range(9):
                for j in range(len(pattern_numero[i])):
                    if re.match(pattern_numero[i][j], paragraph.text):
                        numero_is_inside = True
                        break
        for paragraph in file.paragraphs:
            if numero_is_inside:
                for i in range(9):
                        for j in range(len(pattern_numero[i])):
                            if re.match(pattern_numero[i][j], paragraph.text):
                                if j == 0 or j == 1:
                                    pattern_Debut.append(re.sub(r'^\d+\)|^\d+/', '', paragraph.text))
                                else:
                                    pattern_Debut.append(pattern_numero[i][j])
                                    #pattern_Debut.append(paragraph.text)
            else :
                for i in range(len(pattern_point)):
                    if re.search(pattern_point[i], paragraph.text):
                        pattern_Debut.append(paragraph.text)

    #pattern_Debut_clean = [re.sub(r'^\d+\)|^\d+/|^Q\d+\s*\s*', '', list) for list in pattern_Debut]
        #Remove 1) or 
    #pattern_Debut.append(re.sub(r'^Q\d+\s*\s*', '', paragraph.text))
                        
    for i in range(1,len(pattern_Debut)):
        pattern_End.append(pattern_Debut[i])
    pattern_End.append("///")

    #print("hello")
    
    return pattern_Debut, pattern_End
# --------------------------------------------------------------------------------------------------------------