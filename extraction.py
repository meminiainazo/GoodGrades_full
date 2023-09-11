import pandas as pd
import os
import re
from fill_variables import *
from patterns import *
from docx import Document
import PyPDF2
import mysql.connector
import pymysql
import langdetect

#Global file ------------------------------------------------------------------------------------------------------------
template = pd.read_excel("template.xlsx")
grades_Folder = "grades/"
subject_Folder = "subject/"
reference_Folder = "reference_answer/"
students_answer_Folder = "students_answer/"


grade = get_file(grades_Folder)
subject = get_file(subject_Folder)
reference = get_file(reference_Folder)
# -----------------------------------------------------------------------------------------------------------------------

#Config mysql -----------------------------------------------------------------------------------------------------------
config = {
    'user': 'goodgrades_user',
    'password': 'goodgrades',
    'host': '127.0.0.1',
    'database': 'goodgrades_db',
}
# -----------------------------------------------------------------------------------------------------------------------

#cnx = mysql.connector.connect(**config)
#cur = cnx.cursor(buffered=True)
#Connexion establishment ------------------------------------------------------------------------------------------------
cnx = pymysql.connect(**config)
cur = cnx.cursor()
# -----------------------------------------------------------------------------------------------------------------------

#Create table if not exist ----------------------------------------------------------------------------------------------
cur.execute("""CREATE TABLE IF NOT EXISTS goodgrades (
id VARCHAR(30),
utilisateur VARCHAR(70),
ecole VARCHAR(70),
langue VARCHAR(70),
annees VARCHAR(20),
session VARCHAR(100),
competence VARCHAR(100),
diplome VARCHAR(100),
question_type INT,
question_numero INT,
support TEXT,
enonce TEXT,
reponse_referente TEXT,
reponse_apprenant TEXT,
note_sur_10 FLOAT,
note FLOAT,
note_sur FLOAT,
copie LONGBLOB,
nom_copie VARCHAR(200))"""
)
# -----------------------------------------------------------------------------------------------------------------------

#Get question grade ------------------------------------------------------------------------------------------------------
def get_question_grade(grade, grade_pattern):
    target_columns = []
    for column in grade.columns.tolist():
        if re.search(grade_pattern, str(grade.at[0, column])):
            target_columns.append(column)
    return target_columns
# ------------------------------------------------------------------------------------------------------------------------

#Extract bareme ---------------------------------------------------------------------------------------------------------
def extract_bareme(column):
    max_Grade_pattern = r"Max grade:\s*(\d+)"
    if re.search(max_Grade_pattern, column):
        match = re.search(max_Grade_pattern, column)
        return float(match.group(1))
# ------------------------------------------------------------------------------------------------------------------------

#Extract questions -------------------------------------------------------------------------------------------------------
def extract_between(file, start_pattern, end_pattern):
    extracted_paragraphs = ""
    if file.endswith(".docx"):
        file = Document(file)
        is_inside_paragraph = False
        for paragraph in file.paragraphs:
            if re.search(end_pattern, paragraph.text, re.I):
                break
            if re.search(start_pattern, paragraph.text, re.I):
                is_inside_paragraph = True
            if is_inside_paragraph:
                extracted_paragraphs += ("\n" + paragraph.text)
        #print("non-docx :", file)
    elif file.endswith(".pdf"):
        text = ""
        file = PyPDF2.PdfReader(file)
        for num_page in range(len(file.pages)):
            page = file.pages[num_page]
            text += page.extract_text()
        #print(type(text))
        result = re.findall(f"{start_pattern}(.*?){end_pattern}", text, re.DOTALL)
        # Check if the patterns were found and extract the text
        if result:
            extracted_paragraphs = result[0]
            print(extracted_paragraphs)
    return extracted_paragraphs
# ------------------------------------------------------------------------------------------------------------------------

#Insert data into the database -------------------------------------------------------------------------------------------
def insert_data(table_name, data_dict):
    # Generate the SQL statement dynamically
    columns = ', '.join(data_dict.keys())
    placeholders = ', '.join(['%s' for _ in data_dict.values()])
    insert_data_sql = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"
    
    # Extract the values from the data_dict
    values = tuple(data_dict.values())
    
    # Execute the SQL statement with the values
    cur.execute(insert_data_sql, values)
    cnx.commit()
# ------------------------------------------------------------------------------------------------------------------------

#Insert file -------------------------------------------------------------------------------------------------------------
def insert_file(file):
    with open(file, "rb") as file:
        file_content = file.read()
    return file_content
# ------------------------------------------------------------------------------------------------------------------------

#Extract name and grades -------------------------------------------------------------------------------------------------
def extract_all(file, question_type, question_numero, ecole, langue, annees, session, competence, bareme, support, enonce,reponse, reponse_referente, name_index, column, grade_pattern, copie, nom_copie):
    if re.search(grade_pattern, str(file.at[name_index, column])):
        match = re.search(grade_pattern, str(file.at[name_index, column]))
        data = {
            "utilisateur" : file.iloc[name_index, 0],
            "ecole" : ecole,
            "langue" : langue,
            "annees" : annees,
            "session" : session,
            "competence" : competence,
            "note" : float(match.group(1)),
            "note_sur" : bareme,
            "note_sur_10" : round(float(match.group(1))*10/bareme, 2),
            "support" : support,
            "enonce" : enonce,
            "question_type" : question_type,
            "question_numero" : question_numero,
            "reponse_referente" : reponse_referente,
            "reponse_apprenant" : reponse,
            "copie" : copie,
            "nom_copie" : nom_copie
        }
        insert_data("goodgrades", data)
        # template.at[j, "note"] = float(match.group(1))
        # template.at[j, "note_sur"] = bareme
        # template.at[j, "utilisateur"] = file.iloc[name_index, 0]
        #->template.at[j, "note_sur_10"] = round(template.at[j, "note"]*10/template.at[j, "note_sur"], 2)
    else :
        data = {
            "utilisateur" : file.iloc[name_index, 0],
            "ecole" : ecole,
            "langue" : langue,
            "annees" : annees,
            "session" : session,
            "competence" : competence,
            "note" : float(file.at[name_index, column]),
            "note_sur" : bareme,
            "note_sur_10" : round(float(match.group(1))*10/bareme, 2),
            "support" : support,
            "enonce" : enonce,
            "question_type" : question_type,
            "question_numero" : question_numero,
            "reponse_referente" : reponse_referente,
            "reponse_apprenant" : reponse,
            "copie" : copie,
            "nom_copie" : nom_copie
        }
        insert_data("goodgrades", data)
        # template.at[j, "note"] = float(file.at[name_index, column])
        # template.at[j, "note_sur"] = bareme
        # template.at[j, "utilisateur"] = file.iloc[name_index, 0]
        #->template.at[j, "note_sur_10"] = round(template.at[j, "note"]*10/template.at[j, "note_sur"], 2)
#-------------------------------------------------------------------------------------------------------------------------

#Extraction support ------------------------------------------------------------------------------------------------------
def extract_support(file):
    support = ""
    if file.endswith(".docx"):
        doc = Document(file)
        for paragraph in doc.paragraphs:
            support += paragraph.text
    elif file.endswith(".pdf"):
        doc = PyPDF2.PdfReader(file)
        for page_num in range(doc.numPages):
            page = doc.getPage(page_num)
            support += page.extract_text()
    return support
# ------------------------------------------------------------------------------------------------------------------------

#Extraction --------------------------------------------------------------------------------------------------------------
def run_extraction(grade):
    file = pd.read_excel(grade)
    grade_pattern = r"Grade:\s*(\d+\.?\d*)"
    target_columns = get_question_grade(file, grade_pattern) #Take the list of cells who contains the grades 
    start_pattern, end_pattern = get_patterns(subject_Folder + subject) #Take the patterns
    ecole = input("Entrez la nom de l'établissement :")
    annees = input("Année scolaire :")
    session = input("Session :")
    competence = input("Examen :")
    j,k = 0,0 #For looping the row and filling baremes
    cur.execute("TRUNCATE TABLE goodgrades")
    for name_index in range(file.shape[0]):
        for column in target_columns:
            answer_Folder = students_answer_Folder + file.iloc[name_index][0] + "/"
            #there_is_Table_or_image = check_table_or_image(answer_Folder + get_file(answer_Folder))
            extract_all(file,
                        1 if len(extract_between(reference_Folder + reference, start_pattern[k], end_pattern[k]))>300 else 2,
                        k + 1,
                        ecole,
                        langdetect.detect(extract_between(reference_Folder + reference, start_pattern[k], end_pattern[k])).upper(),
                        annees,
                        session,
                        competence,
                        extract_bareme(column),
                        extract_support(subject_Folder + subject),
                        extract_between(subject_Folder + subject, start_pattern[k], end_pattern[k]), 
                        extract_between(answer_Folder + get_file(answer_Folder), start_pattern[k], end_pattern[k]),
                        extract_between(reference_Folder + reference, start_pattern[k], end_pattern[k]),
                        name_index,
                        column, 
                        grade_pattern,
                        insert_file(answer_Folder + get_file(answer_Folder)),
                        get_file(answer_Folder)
                       )

            # template.at[j, "enonce"] = extract_between(subject_Folder + subject, start_pattern[k], end_pattern[k]) #Extract questions
            # template.at[j, "reponse_apprenant"] = extract_between(answer_Folder + get_file(answer_Folder), start_pattern[k], end_pattern[k])
            if k>len(target_columns)-2:
                k = 0
            else :
                k += 1
            # j += 1
    #template.to_excel("tmplt.xlsx")
# ------------------------------------------------------------------------------------------------------------------------